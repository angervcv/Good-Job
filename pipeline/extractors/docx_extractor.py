"""
DOCX 文本提取器 (基于 python-docx)
"""

from pathlib import Path
from docx import Document
import config


def extract_docx_text(file_path: str | Path) -> list[dict]:
    """提取 DOCX 中每个段落的文本

    Returns:
        [{"index": int, "style": str, "text": str}, ...]
    """
    file_path = Path(file_path)
    doc = Document(str(file_path))
    paragraphs = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:  # 只保留非空段落
            paragraphs.append({
                "index": i,
                "style": para.style.name if para.style else "Normal",
                "text": text,
            })

    return paragraphs


def extract_docx_full_text(file_path: str | Path) -> str:
    """提取 DOCX 全部文本"""
    file_path = Path(file_path)
    doc = Document(str(file_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_docx_with_tables(file_path: str | Path) -> str:
    """提取 DOCX 文本+表格内容"""
    file_path = Path(file_path)
    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for i, table in enumerate(doc.tables):
        parts.append(f"\n[表格 {i + 1}]")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def scan_docx_files() -> list[dict]:
    """扫描题目目录下所有 DOCX 文件"""
    questions_dir = config.QUESTIONS_DIR
    files = []

    for docx_path in questions_dir.rglob("*.docx"):
        try:
            paras = extract_docx_text(docx_path)
            total_chars = sum(len(p["text"]) for p in paras)
            files.append({
                "path": str(docx_path),
                "relative_path": str(docx_path.relative_to(questions_dir)),
                "paragraphs": len(paras),
                "total_chars": total_chars,
            })
        except Exception as e:
            print(f"  [警告] 无法读取 {docx_path}: {e}")

    return files


if __name__ == "__main__":
    import sys
    sys.path.insert(0, str(config.ROOT_DIR))

    files = scan_docx_files()
    for f in files:
        print(f"{f['relative_path']} ({f['paragraphs']}段, {f['total_chars']}字符)")